"""
Word文档解析器
"""
from docx import Document
from typing import List, Dict


class WordParser:
    """解析Word文档(.docx)"""

    @staticmethod
    def parse(file_path: str) -> Dict:
        doc = Document(file_path)
        result = {
            'text': [],
            'tables': [],
            'headings': [],
            'paragraphs': []
        }

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                result['paragraphs'].append(text)
                if paragraph.style.name.startswith('Heading'):
                    result['headings'].append({
                        'level': int(paragraph.style.name.replace('Heading ', '')) if 'Heading' in paragraph.style.name else 0,
                        'text': text
                    })
                else:
                    result['text'].append(text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result['tables'].append(table_data)

        return result

    @staticmethod
    def extract_requirements(paragraphs: List[str]) -> List[Dict]:
        requirements = []
        keywords = ['需求', '功能', '模块', '应该', '需要', '要求', '支持', '提供']
        for para in paragraphs:
            for keyword in keywords:
                if keyword in para:
                    requirements.append({
                        'type': 'functional',
                        'description': para,
                        'keyword': keyword
                    })
                    break
        return requirements

    @staticmethod
    def extract_api_info(paragraphs: List[str]) -> List[Dict]:
        apis = []
        api_keywords = ['接口', 'API', '请求', '响应', 'URL', 'POST', 'GET', 'PUT', 'DELETE']
        for para in paragraphs:
            for keyword in api_keywords:
                if keyword in para.upper() if keyword == 'API' else keyword in para:
                    apis.append({'description': para, 'detected_keyword': keyword})
                    break
        return apis
